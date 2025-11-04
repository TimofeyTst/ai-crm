"""Resume document generator service (DOCX/PDF)."""

from datetime import date
import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)

from ai_crm.pkg import context
from ai_crm.pkg.logger import logger as logger_lib
from ai_crm.pkg.models.ai_crm import parsed_resume as parsed_resume_models
from ai_crm.pkg.models.exceptions import ai as ai_exceptions

logger = logger_lib.get_logger(__name__)


async def generate_docx_resume(
    context: context.AnyContext,
    parsed_resume: parsed_resume_models.ParsedResume,
) -> bytes:
    """Generate DOCX resume from parsed data."""
    try:
        logger.info(f"Generating DOCX resume for: {parsed_resume.name}")

        doc = Document()
        _set_docx_document_margins(doc)
        _add_docx_header(doc, parsed_resume)
        _add_docx_summary(doc, parsed_resume)
        _add_docx_professional_experience(doc, parsed_resume)
        _add_docx_education(doc, parsed_resume)
        _add_docx_certifications(doc, parsed_resume)
        _add_docx_skills(doc, parsed_resume)

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        logger.info("Successfully generated DOCX resume")
        return output.getvalue()

    except Exception as e:
        logger.exception(f"DOCX generation failed: {e}")
        raise ai_exceptions.ResumeGenerationFailed from e


async def generate_pdf_resume(
    context: context.AnyContext,
    parsed_resume: parsed_resume_models.ParsedResume,
) -> bytes:
    """Generate PDF resume from parsed data."""
    try:
        logger.info(f"Generating PDF resume for: {parsed_resume.name}")

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=0.75 * inch,
            leftMargin=0.75 * inch,
            topMargin=0.5 * inch,
            bottomMargin=0.5 * inch,
        )

        styles = _get_pdf_styles()
        story = []

        _add_pdf_header(story, styles, parsed_resume)
        _add_pdf_summary(story, styles, parsed_resume)
        _add_pdf_professional_experience(story, styles, parsed_resume)
        _add_pdf_education(story, styles, parsed_resume)
        _add_pdf_certifications(story, styles, parsed_resume)
        _add_pdf_skills(story, styles, parsed_resume)

        doc.build(story)
        buffer.seek(0)

        logger.info("Successfully generated PDF resume")
        return buffer.getvalue()

    except Exception as e:
        logger.exception(f"PDF generation failed: {e}")
        raise ai_exceptions.ResumeGenerationFailed from e


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    """Add hyperlink to paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)

    # Установка цвета (синий)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")  # синий цвет
    rPr.append(color)

    # Подчеркивание
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    # Размер
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "24")  # 12 pt
    rPr.append(size)

    sizeCs = OxmlElement("w:szCs")
    sizeCs.set(qn("w:val"), "24")  # 12 pt
    rPr.append(sizeCs)

    new_run.append(rPr)

    new_text = OxmlElement("w:t")
    new_text.text = text
    new_run.append(new_text)

    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)


def _set_docx_document_margins(doc: Document) -> None:
    """Set DOCX document margins."""
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.44)
        section.bottom_margin = Inches(0.29)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)


def _add_docx_header(
    doc: Document, resume: parsed_resume_models.ParsedResume
) -> None:
    """Add DOCX header with name and contact."""
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(resume.name)
    name_run.font.name = "Arial"
    name_run.font.size = Pt(20)
    name_run.font.bold = True
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(0)

    if resume.position:
        position_para = doc.add_paragraph()
        position_run = position_para.add_run(resume.position)
        position_run.font.name = "Arial"
        position_run.font.size = Pt(14)
        position_run.font.bold = True
        position_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        position_para.paragraph_format.space_after = Pt(0)

    if resume.email or resume.linkedin:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.paragraph_format.space_after = Pt(0)

        if resume.email:
            _add_hyperlink(contact_para, f"mailto:{resume.email}", resume.email)

        if resume.email and resume.linkedin:
            separator_run = contact_para.add_run(" | ")
            separator_run.font.name = "Arial"

        if resume.linkedin:
            _add_hyperlink(contact_para, resume.linkedin, "LinkedIn")

    doc.add_paragraph()


def _add_docx_summary(
    doc: Document, resume: parsed_resume_models.ParsedResume
) -> None:
    """Add DOCX summary section."""
    if not resume.summary:
        return

    _add_docx_section_header(doc, "SUMMARY")
    summary_para = doc.add_paragraph(resume.summary)
    summary_para.runs[0].font.size = Pt(11)
    summary_para.runs[0].font.name = "Arial"
    summary_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    summary_para.paragraph_format.space_after = Pt(10)
    doc.add_paragraph()


def _clear_table_borders(table) -> None:
    """
    Убирает все границы у указанной таблицы docx, делая ее невидимой.
    """
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'nil')
                tcBorders.append(border)
            tcPr.append(tcBorders)

def _add_docx_professional_experience(
    doc: Document, resume: parsed_resume_models.ParsedResume
) -> None:
    """Add DOCX professional experience."""
    if not resume.professional_experience:
        return

    _add_docx_section_header(doc, "PROFESSIONAL EXPERIENCE")

    for exp in resume.professional_experience:
        table = doc.add_table(rows=1, cols=2) # TODO: date range в самую правую границу не упирается
        _clear_table_borders(table)

        # Левая ячейка - Название компании
        cell_company = table.cell(0, 0)
        para_company = cell_company.paragraphs[0]
        run_company = para_company.add_run(exp.company_name)
        run_company.font.bold = True
        run_company.font.size = Pt(12)
        run_company.font.name = "Arial"
        para_company.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para_company.paragraph_format.space_before = Pt(0)
        para_company.paragraph_format.space_after = Pt(0)

        # Правая ячейка - Диапазон дат
        cell_date = table.cell(0, 1)
        para_date = cell_date.paragraphs[0]
        date_range = _format_date_range(exp.hired_date, exp.fired_date)
        if date_range:
            run_date = para_date.add_run(date_range)
            run_date.font.bold = True
            run_date.font.size = Pt(12)
            run_date.font.name = "Arial"
            para_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            para_date.paragraph_format.space_before = Pt(0)
            para_date.paragraph_format.space_after = Pt(0)

        position_para = doc.add_paragraph()
        position_run = position_para.add_run(exp.position)
        position_run.font.bold = True
        position_run.font.name = "Arial"
        position_para.paragraph_format.space_after = Pt(1)

        if exp.achievements:
            for achievement in exp.achievements:
                ach_para = doc.add_paragraph(achievement, style="List Bullet")
                ach_para.runs[0].font.name = "Arial"
                ach_para.runs[0].font.size = Pt(11)
                ach_para.paragraph_format.space_after = Pt(1)

        doc.add_paragraph()


def _add_docx_education(
    doc: Document, resume: parsed_resume_models.ParsedResume
) -> None:
    """Add DOCX education section."""
    if not resume.education:
        return

    _add_docx_section_header(doc, "EDUCATION")

    # TODO: as bullets
    for edu in resume.education:
        table = doc.add_table(rows=1, cols=2) # TODO: date range в самую правую границу не упирается
        _clear_table_borders(table)

        cell_degree = table.cell(0, 0)
        degree_para = cell_degree.paragraphs[0]
        if edu.degree:
            degree_run = degree_para.add_run(f"{edu.degree}, {edu.institution}")
            degree_run.font.bold = True
            degree_run.font.name = "Arial"
            degree_run.font.size = Pt(12)
            degree_para.paragraph_format.space_after = Pt(0)
        else:
            inst_run = degree_para.add_run(edu.institution)
            inst_run.font.bold = True
            inst_run.font.name = "Arial"
            inst_run.font.size = Pt(12)
            degree_para.paragraph_format.space_after = Pt(0)

        degree_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        degree_para.paragraph_format.space_before = Pt(0)
        degree_para.paragraph_format.space_after = Pt(0)

        cell_date = table.cell(0, 1)
        date_para = cell_date.paragraphs[0]
        date_range = _format_date_range_years(edu.from_date, edu.to_date)
        if date_range:
            date_run = date_para.add_run(date_range)
            date_run.font.size = Pt(12)
            date_run.font.bold = True
            date_run.font.name = "Arial"
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            date_para.paragraph_format.space_before = Pt(0)
            date_para.paragraph_format.space_after = Pt(0)

        if edu.description:
            desc_para = doc.add_paragraph(edu.description)
            desc_para.runs[0].font.size = Pt(11)
            desc_para.runs[0].font.name = "Arial"
            desc_para.paragraph_format.space_after = Pt(12)

        doc.add_paragraph()


def _add_docx_certifications(
    doc: Document, resume: parsed_resume_models.ParsedResume
) -> None:
    """Add DOCX certifications section."""
    if not resume.certifications:
        return

    _add_docx_section_header(doc, "CERTIFICATIONS & ADDITIONAL TRAINING")

    for cert in resume.certifications:
        bullet_text = cert.name
        if cert.description:
            bullet_text = f"{bullet_text} - {cert.description}"

        if cert.date_obtained:
            date_str = cert.date_obtained.strftime("%Y")
            bullet_text += f", {date_str}"

        cert_para = doc.add_paragraph(bullet_text, style="List Bullet")
        cert_para.runs[0].font.size = Pt(11)
        cert_para.runs[0].font.name = "Arial"
        cert_para.paragraph_format.space_after = Pt(0)

    doc.add_paragraph()


def _add_docx_skills(
    doc: Document, resume: parsed_resume_models.ParsedResume
) -> None:
    """Add DOCX skills section."""
    if not any(
        [
            resume.skills.tech_stack,
            resume.skills.languages,
        ]
    ):
        return

    _add_docx_section_header(doc, "SKILLS")

    # TODO: as bullets
    if resume.skills.tech_stack:
        tech_para = doc.add_paragraph()
        tech_para.add_run("TechStack: ").font.bold = True
        tech_para.add_run(resume.skills.tech_stack)
        tech_para.runs[0].font.size = Pt(11)
        tech_para.runs[1].font.size = Pt(11)
        tech_para.runs[0].font.name = "Arial"
        tech_para.runs[1].font.name = "Arial"
        tech_para.paragraph_format.space_after = Pt(0)

    if resume.skills.languages:
        lang_para = doc.add_paragraph()
        lang_para.add_run("Languages: ").font.bold = True
        lang_para.add_run(resume.skills.languages)
        lang_para.runs[0].font.size = Pt(11)
        lang_para.runs[1].font.size = Pt(11)
        lang_para.runs[0].font.name = "Arial"
        lang_para.runs[1].font.name = "Arial"
        lang_para.paragraph_format.space_after = Pt(0)


def _add_docx_section_header(doc: Document, title: str) -> None:
    """Add DOCX section header."""
    header_para = doc.add_paragraph()
    header_run = header_para.add_run(title)
    header_run.font.name = "Arial"
    header_run.font.size = Pt(12)
    header_run.font.bold = True
    header_para.paragraph_format.space_before = Pt(1)
    header_para.paragraph_format.space_after = Pt(1)
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = header_para._p  # XML абзац
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')      # толщина
    bottom.set(qn('w:space'), '1')    # отступ
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _get_pdf_styles():
    """Get PDF paragraph styles."""
    styles = getSampleStyleSheet()

    styles.add(
        ParagraphStyle(
            name="ResumeTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=20,
            textColor=colors.HexColor("#000000"),
            alignment=1,
            spaceAfter=2,
        )
    )

    styles.add(
        ParagraphStyle(
            name="ResumeSubtitle",
            parent=styles["Normal"],
            fontName="Times-Italic",
            fontSize=11,
            textColor=colors.HexColor("#000000"),
            alignment=1,
            spaceAfter=2,
        )
    )

    styles.add(
        ParagraphStyle(
            name="ContactInfo",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=10,
            textColor=colors.HexColor("#000000"),
            alignment=1,
            spaceAfter=12,
        )
    )

    styles.add(
        ParagraphStyle(
            name="SectionHeader",
            parent=styles["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=12,
            textColor=colors.HexColor("#000000"),
            spaceBefore=8,
            spaceAfter=4,
            borderWidth=1,
            borderColor=colors.HexColor("#000000"),
            borderPadding=2,
        )
    )

    styles.add(
        ParagraphStyle(
            name="CompanyName",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=11,
            spaceAfter=0,
        )
    )

    styles.add(
        ParagraphStyle(
            name="Position",
            parent=styles["Normal"],
            fontName="Times-Italic",
            fontSize=10,
            spaceAfter=0,
        )
    )

    styles.add(
        ParagraphStyle(
            name="DateRange",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=9,
            spaceAfter=4,
        )
    )

    styles.add(
        ParagraphStyle(
            name="BulletItem",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=10,
            leftIndent=20,
            spaceAfter=2,
        )
    )

    return styles


def _add_pdf_header(story, styles, resume: parsed_resume_models.ParsedResume):
    """Add PDF header."""
    story.append(Paragraph(resume.name, styles["ResumeTitle"]))

    if resume.position:
        story.append(Paragraph(resume.position, styles["ResumeSubtitle"]))

    contact_parts = []
    if resume.email:
        contact_parts.append(resume.email)
    if resume.linkedin:
        contact_parts.append("LinkedIn")

    if contact_parts:
        contact_text = " | ".join(contact_parts)
        story.append(Paragraph(contact_text, styles["ContactInfo"]))

    story.append(Spacer(1, 0.1 * inch))


def _add_pdf_summary(story, styles, resume: parsed_resume_models.ParsedResume):
    """Add PDF summary."""
    if not resume.summary:
        return

    story.append(Paragraph("SUMMARY", styles["SectionHeader"]))
    story.append(Paragraph(resume.summary, styles["Normal"]))
    story.append(Spacer(1, 0.15 * inch))


def _add_pdf_professional_experience(
    story, styles, resume: parsed_resume_models.ParsedResume
):
    """Add PDF professional experience."""
    if not resume.professional_experience:
        return

    story.append(Paragraph("PROFESSIONAL EXPERIENCE", styles["SectionHeader"]))

    for exp in resume.professional_experience:
        story.append(Paragraph(exp.company_name, styles["CompanyName"]))
        story.append(Paragraph(exp.position, styles["Position"]))

        date_range = _format_date_range(exp.hired_date, exp.fired_date)
        if date_range:
            story.append(Paragraph(date_range, styles["DateRange"]))

        if exp.achievements:
            for achievement in exp.achievements:
                bullet_text = f"• {achievement}"
                story.append(Paragraph(bullet_text, styles["BulletItem"]))

        story.append(Spacer(1, 0.1 * inch))


def _add_pdf_education(
    story, styles, resume: parsed_resume_models.ParsedResume
):
    """Add PDF education."""
    if not resume.education:
        return

    story.append(Paragraph("EDUCATION", styles["SectionHeader"]))

    for edu in resume.education:
        if edu.degree:
            text = f"{edu.degree}, {edu.institution}"
        else:
            text = edu.institution

        story.append(Paragraph(text, styles["CompanyName"]))

        date_range = _format_date_range(edu.from_date, edu.to_date)
        if date_range:
            story.append(Paragraph(date_range, styles["DateRange"]))

        if edu.description:
            story.append(Paragraph(edu.description, styles["Normal"]))

        story.append(Spacer(1, 0.1 * inch))


def _add_pdf_certifications(
    story, styles, resume: parsed_resume_models.ParsedResume
):
    """Add PDF certifications."""
    if not resume.certifications:
        return

    story.append(
        Paragraph(
            "CERTIFICATIONS & ADDITIONAL TRAINING", styles["SectionHeader"]
        )
    )

    for cert in resume.certifications:
        bullet_text = f"• {cert.name}"
        if cert.date_obtained:
            date_str = cert.date_obtained.strftime("%Y")
            bullet_text += f", {date_str}"
        story.append(Paragraph(bullet_text, styles["BulletItem"]))

    story.append(Spacer(1, 0.1 * inch))


def _add_pdf_skills(story, styles, resume: parsed_resume_models.ParsedResume):
    """Add PDF skills."""
    if not any([resume.skills.tech_stack, resume.skills.languages]):
        return

    story.append(Paragraph("SKILLS", styles["SectionHeader"]))

    if resume.skills.tech_stack:
        text = f"<b>TechStack:</b> {resume.skills.tech_stack}"
        story.append(Paragraph(text, styles["Normal"]))

    if resume.skills.languages:
        text = f"<b>Languages:</b> {resume.skills.languages}"
        story.append(Paragraph(text, styles["Normal"]))


def _format_date_range(from_date: date | None, to_date: date | None) -> str:
    """Format date range."""
    if not from_date:
        return ""

    from_str = from_date.strftime("%b %Y")
    to_str = "till now" if to_date is None else to_date.strftime("%b %Y")

    return f"{from_str} – {to_str}"


def _format_date_range_years(from_date: date | None, to_date: date | None) -> str:
    """Format date range."""
    if not from_date:
        return ""

    from_str = from_date.strftime("%Y")
    to_str = "till now" if to_date is None else to_date.strftime("%Y")

    return f"{from_str} – {to_str}"
